from docx import Document
from docx.shared import Inches
from typing import List
from document_loader import ExtractedImage


class ReportBuilder:
    def _parse_image_id(self, img_id):
        """Parse image ID, handling both int and 'Image X' formats."""
        if isinstance(img_id, int):
            return img_id
        if isinstance(img_id, str):
            if img_id.startswith("Image "):
                try:
                    return int(img_id.split(" ")[1])
                except (ValueError, IndexError):
                    return None
            try:
                return int(img_id)
            except ValueError:
                return None
        return None

    def build(self, report_json, output_path: str, extracted_images: List[ExtractedImage] = None):
        extracted_images = extracted_images or []
        doc = Document()
        title = report_json.get("title", "Generated Report")
        doc.add_heading(title, level=0)

        sections = report_json.get("sections", [])
        for section in sections:
            heading = section.get("heading")
            content = section.get("content", "")
            image_ids = section.get("images", [])

            if heading:
                doc.add_heading(heading, level=1)

            if isinstance(content, list):
                for paragraph in content:
                    doc.add_paragraph(str(paragraph))
            else:
                for paragraph in str(content).split("\n\n"):
                    text = paragraph.strip()
                    if text:
                        doc.add_paragraph(text)

            # Add images for this section
            if image_ids and extracted_images:
                for img_id_raw in image_ids:
                    img_id = self._parse_image_id(img_id_raw)
                    if img_id is not None and 0 <= img_id < len(extracted_images):
                        img = extracted_images[img_id]
                        import io
                        img_buffer = io.BytesIO()
                        img.image.save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        
                        doc.add_picture(img_buffer, width=Inches(4))
                        # Add a caption
                        doc.add_paragraph(f"Figure: {img.get_description()}")

        doc.save(output_path)
